import PyPDF2
import docx
import email
import io
from typing import List, Dict, Any, Tuple
from pathlib import Path

class DocumentLoader:
    """Handles loading and parsing of various document types"""
    
    @staticmethod
    def load_pdf(file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Load and extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            metadata = {
                "total_pages": len(pdf_reader.pages),
                "page_texts": {}
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                metadata["page_texts"][page_num + 1] = page_text
            
            return text_content.strip(), metadata
            
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def load_docx(file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Load and extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text_content = ""
            metadata = {
                "total_paragraphs": len(doc.paragraphs),
                "paragraph_texts": {}
            }
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content += f"{paragraph.text}\n"
                    metadata["paragraph_texts"][i] = paragraph.text
            
            return text_content.strip(), metadata
            
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def load_email(file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Load and extract text from email file"""
        try:
            email_content = file_content.decode('utf-8')
            msg = email.message_from_string(email_content)
            
            # Extract basic email metadata
            metadata = {
                "subject": msg.get("Subject", ""),
                "from": msg.get("From", ""),
                "to": msg.get("To", ""),
                "date": msg.get("Date", ""),
            }
            
            # Extract email body
            text_content = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text_content += part.get_payload(decode=True).decode('utf-8')
            else:
                text_content = msg.get_payload(decode=True).decode('utf-8')
            
            # Format the email content
            formatted_content = f"""
Subject: {metadata['subject']}
From: {metadata['from']}
To: {metadata['to']}
Date: {metadata['date']}

{text_content}
            """.strip()
            
            return formatted_content, metadata
            
        except Exception as e:
            raise ValueError(f"Error parsing email: {str(e)}")
    
    @classmethod
    def load_document(cls, file_content: bytes, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """Main method to load document based on file type"""
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return cls.load_pdf(file_content)
        elif file_type in ["docx", "doc"]:
            return cls.load_docx(file_content)
        elif file_type in ["eml", "email"]:
            return cls.load_email(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def preprocess_text(text: str) -> str:
        """Clean and preprocess extracted text"""
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove special characters that might interfere with processing
        # Keep basic punctuation
        import re
        text = re.sub(r'[^\w\s.,!?;:()\-"]', ' ', text)
        
        return text.strip()